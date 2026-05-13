#!/usr/bin/env python3
"""
Comprehensive Tesseract OCR Script
Supports single image, batch processing, multiple output formats (including .docx), and preprocessing
"""

import os
import sys
import argparse
import json
import csv
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import concurrent.futures
from dataclasses import dataclass, asdict
from enum import Enum

# Third-party imports
try:
    import pytesseract
    from PIL import Image, ImageEnhance, ImageFilter
    import cv2
    import numpy as np
    from docx import Document
    from docx.shared import Inches
except ImportError as e:
    print(f"Error: Missing required library. Install with: pip install pytesseract pillow opencv-python numpy python-docx")
    print(f"Details: {e}")
    sys.exit(1)

# Path to Tesseract installation
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# ============================================================================
# Configuration Classes
# ============================================================================

class PSMMode(Enum):
    """Page Segmentation Modes for Tesseract"""
    OSD_ONLY = 0
    AUTO_OSD = 1
    AUTO_ONLY = 3
    SINGLE_COLUMN = 4
    VERTICAL_TEXT = 5
    UNIFORM_BLOCK = 6
    SINGLE_LINE = 7
    SINGLE_WORD = 8
    SINGLE_WORD_CIRCLE = 9
    SPARSE_TEXT = 11
    SPARSE_TEXT_OSD = 12
    RAW_LINE = 13

class OEMMode(Enum):
    """OCR Engine Modes"""
    LEGACY_ONLY = 0
    LSTM_ONLY = 1
    LEGACY_LSTM = 2
    DEFAULT = 3

@dataclass
class OCRConfig:
    """Configuration for OCR processing"""
    language: str = 'eng'
    psm: PSMMode = PSMMode.UNIFORM_BLOCK
    oem: OEMMode = OEMMode.DEFAULT
    tessdata_dir: Optional[str] = None
    tesseract_path: Optional[str] = None
    confidence_threshold: float = 0.0
    preserve_layout: bool = True
    output_format: str = 'txt'  # txt, json, csv, hocr, pdf, docx

@dataclass
class PreprocessConfig:
    """Configuration for image preprocessing"""
    grayscale: bool = True
    threshold: bool = True
    threshold_value: int = 150
    denoise: bool = True
    denoise_strength: int = 3
    resize: bool = False
    resize_factor: float = 2.0
    sharpen: bool = False
    contrast_enhance: bool = False
    contrast_factor: float = 1.5
    deskew: bool = True
    remove_borders: bool = True

# ============================================================================
# Image Preprocessor
# ============================================================================

class ImagePreprocessor:
    """Handles image preprocessing to improve OCR accuracy"""
    
    def __init__(self, config: PreprocessConfig):
        self.config = config
    
    def preprocess(self, image_path: str) -> np.ndarray:
        """Apply preprocessing pipeline to image"""
        # Read image
        img = cv2.imread(image_path)
        if img is None:
            raise ValueError(f"Cannot read image: {image_path}")
        
        # Convert to grayscale
        if self.config.grayscale:
            img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Deskew
        if self.config.deskew:
            img = self._deskew(img)
        
        # Remove borders
        if self.config.remove_borders:
            img = self._remove_borders(img)
        
        # Apply thresholding
        if self.config.threshold:
            _, img = cv2.threshold(img, self.config.threshold_value, 255, cv2.THRESH_BINARY)
        
        # Denoise
        if self.config.denoise:
            img = cv2.medianBlur(img, self.config.denoise_strength)
        
        # Resize (upscale for better OCR on small text)
        if self.config.resize:
            height, width = img.shape[:2]
            new_size = (int(width * self.config.resize_factor), 
                       int(height * self.config.resize_factor))
            img = cv2.resize(img, new_size, interpolation=cv2.INTER_CUBIC)
        
        # Contrast enhancement
        if self.config.contrast_enhance:
            pil_img = Image.fromarray(img)
            enhancer = ImageEnhance.Contrast(pil_img)
            pil_img = enhancer.enhance(self.config.contrast_factor)
            img = np.array(pil_img)
        
        # Sharpen
        if self.config.sharpen:
            kernel = np.array([[-1,-1,-1],
                              [-1, 9,-1],
                              [-1,-1,-1]])
            img = cv2.filter2D(img, -1, kernel)
        
        return img
    
    def _deskew(self, image: np.ndarray) -> np.ndarray:
        """Deskew image to correct rotation"""
        coords = np.column_stack(np.where(image > 0))
        if len(coords) < 100:
            return image
        
        angle = cv2.minAreaRect(coords)[-1]
        if angle < -45:
            angle = -(90 + angle)
        else:
            angle = -angle
        
        if abs(angle) > 0.5:
            (h, w) = image.shape[:2]
            center = (w // 2, h // 2)
            M = cv2.getRotationMatrix2D(center, angle, 1.0)
            image = cv2.warpAffine(image, M, (w, h), 
                                   flags=cv2.INTER_CUBIC,
                                   borderMode=cv2.BORDER_REPLICATE)
        return image
    
    def _remove_borders(self, image: np.ndarray, border_percent: int = 2) -> np.ndarray:
        """Remove borders from image"""
        h, w = image.shape[:2]
        border_h = int(h * border_percent / 100)
        border_w = int(w * border_percent / 100)
        
        if border_h > 0 and border_w > 0:
            image = image[border_h:-border_h, border_w:-border_w]
        return image
    
    def save_debug_image(self, image: np.ndarray, output_path: str) -> None:
        """Save preprocessed image for debugging"""
        cv2.imwrite(output_path, image)

# ============================================================================
# OCR Engine
# ============================================================================

class OCREngine:
    """Main OCR engine using Tesseract"""
    
    def __init__(self, config: OCRConfig):
        self.config = config
        self._setup_tesseract()
    
    def _setup_tesseract(self):
        """Configure Tesseract path if provided"""
        if self.config.tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = self.config.tesseract_path
    
    def _build_config_string(self) -> str:
        """Build Tesseract configuration string"""
        config_parts = [
            f'--oem {self.config.oem.value}',
            f'--psm {self.config.psm.value}'
        ]
        
        if self.config.tessdata_dir:
            config_parts.append(f'--tessdata-dir "{self.config.tessdata_dir}"')
        
        if self.config.preserve_layout:
            config_parts.append('-c preserve_interword_spaces=1')
        
        return ' '.join(config_parts)
    
    def extract_text(self, image: np.ndarray) -> Tuple[str, Dict]:
        """Extract text from preprocessed image"""
        config = self._build_config_string()
        pil_image = Image.fromarray(image)
        
        try:
            # Get raw text
            text = pytesseract.image_to_string(pil_image, 
                                              lang=self.config.language,
                                              config=config)
            
            # Get detailed data
            data = pytesseract.image_to_data(pil_image,
                                            lang=self.config.language,
                                            config=config,
                                            output_type=pytesseract.Output.DICT)
            
            # Filter by confidence if threshold > 0
            if self.config.confidence_threshold > 0:
                filtered_words = []
                for i, conf in enumerate(data['conf']):
                    if conf >= self.config.confidence_threshold:
                        filtered_words.append(data['text'][i])
                    else:
                        filtered_words.append('')
                filtered_text = ' '.join(filtered_words)
            else:
                filtered_text = text
            
            metadata = {
                'confidence_scores': data['conf'],
                'word_count': len([w for w in data['text'] if w.strip()]),
                'avg_confidence': float(np.mean([c for c in data['conf'] if c > 0])) if data['conf'] else 0
            }
            
            return filtered_text.strip(), metadata
            
        except Exception as e:
            raise RuntimeError(f"OCR processing failed: {e}")
    
    def get_hocr(self, image: np.ndarray) -> str:
        """Get HOCR output (includes positional data)"""
        config = self._build_config_string()
        pil_image = Image.fromarray(image)
        return pytesseract.image_to_pdf_or_hocr(pil_image,
                                               lang=self.config.language,
                                               config=config,
                                               extension='hocr')
    
    def get_pdf(self, image: np.ndarray) -> bytes:
        """Get PDF output with selectable text"""
        config = self._build_config_string()
        pil_image = Image.fromarray(image)
        return pytesseract.image_to_pdf_or_hocr(pil_image,
                                               lang=self.config.language,
                                               config=config,
                                               extension='pdf')

# ============================================================================
# Output Handlers
# ============================================================================

class OutputHandler:
    """Handles saving OCR results in various formats"""
    
    @staticmethod
    def save_txt(text: str, output_path: str, metadata: Dict = None) -> None:
        """Save as plain text"""
        with open(output_path, 'w', encoding='utf-8') as f:
            if metadata:
                f.write(f"# OCR Results\n")
                f.write(f"# Generated: {datetime.now().isoformat()}\n")
                f.write(f"# Average Confidence: {metadata.get('avg_confidence', 0):.2f}%\n")
                f.write(f"# Word Count: {metadata.get('word_count', 0)}\n")
                f.write("\n" + "="*60 + "\n\n")
            f.write(text)
    
    @staticmethod
    def save_docx(text: str, output_path: str, metadata: Dict = None, 
                  original_image: str = None) -> None:
        """Save as Word document (.docx)"""
        doc = Document()
        # Add title
        doc.add_heading('OCR Results', 0)
        # Add metadata
        if metadata:
            doc.add_paragraph(f"Source image: {original_image if original_image else 'Unknown'}")
            doc.add_paragraph(f"Generated: {datetime.now().isoformat()}")
            doc.add_paragraph(f"Average confidence: {metadata.get('avg_confidence', 0):.2f}%")
            doc.add_paragraph(f"Word count: {metadata.get('word_count', 0)}")
            doc.add_paragraph()
        # Add extracted text
        doc.add_heading('Extracted Text', level=1)
        # Split text into paragraphs
        for para in text.split('\n'):
            if para.strip():
                doc.add_paragraph(para)
            else:
                doc.add_paragraph()  # empty line
        # Save
        doc.save(output_path)
    
    @staticmethod
    def save_json(text: str, output_path: str, metadata: Dict = None, 
                  original_image: str = None) -> None:
        """Save as JSON with metadata"""
        result = {
            'file': original_image,
            'timestamp': datetime.now().isoformat(),
            'text': text,
            'metadata': metadata or {}
        }
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)
    
    @staticmethod
    def save_csv(results: List[Dict], output_path: str) -> None:
        """Save multiple results as CSV"""
        if not results:
            return
        
        with open(output_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=['file', 'word_count', 'avg_confidence', 'text_preview'])
            writer.writeheader()
            for r in results:
                writer.writerow({
                    'file': r.get('file', ''),
                    'word_count': r.get('metadata', {}).get('word_count', 0),
                    'avg_confidence': r.get('metadata', {}).get('avg_confidence', 0),
                    'text_preview': r.get('text', '')[:200] + '...' if len(r.get('text', '')) > 200 else r.get('text', '')
                })

# ============================================================================
# Batch Processor
# ============================================================================

class BatchOCRProcessor:
    """Processes multiple images in parallel"""
    
    def __init__(self, ocr_engine: OCREngine, preprocessor: ImagePreprocessor,
                 output_handler: OutputHandler, max_workers: int = 4):
        self.ocr_engine = ocr_engine
        self.preprocessor = preprocessor
        self.output_handler = output_handler
        self.max_workers = max_workers
    
    def process_single(self, image_path: str, output_dir: str, 
                       output_format: str, save_debug: bool = False) -> Dict:
        """Process a single image"""
        try:
            # Preprocess
            processed = self.preprocessor.preprocess(image_path)
            
            # Save debug image if requested
            if save_debug:
                debug_path = os.path.join(output_dir, f"debug_{Path(image_path).stem}.png")
                self.preprocessor.save_debug_image(processed, debug_path)
            
            # Extract text
            text, metadata = self.ocr_engine.extract_text(processed)
            
            return {
                'file': image_path,
                'text': text,
                'metadata': metadata,
                'success': True,
                'error': None
            }
            
        except Exception as e:
            return {
                'file': image_path,
                'text': '',
                'metadata': {},
                'success': False,
                'error': str(e)
            }
    
    def process_batch(self, image_paths: List[str], output_dir: str,
                     output_format: str, save_debug: bool = False) -> List[Dict]:
        """Process multiple images in parallel"""
        results = []
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            futures = {
                executor.submit(self.process_single, path, output_dir, output_format, save_debug): path
                for path in image_paths
            }
            
            for future in concurrent.futures.as_completed(futures):
                result = future.result()
                results.append(result)
                
                # Save individual result - use image stem for output filename
                base_name = Path(result['file']).stem
                output_file = os.path.join(output_dir, f"{base_name}.{output_format}")
                
                if result['success']:
                    if output_format == 'txt':
                        self.output_handler.save_txt(result['text'], output_file, result['metadata'])
                    elif output_format == 'docx':
                        self.output_handler.save_docx(result['text'], output_file, result['metadata'], result['file'])
                    elif output_format == 'json':
                        self.output_handler.save_json(result['text'], output_file, 
                                                      result['metadata'], result['file'])
                    print(f"✓ Processed: {Path(result['file']).name} -> {Path(output_file).name}")
                else:
                    print(f"✗ Failed: {Path(result['file']).name} - {result['error']}")
        
        return results

# ============================================================================
# Main CLI Application
# ============================================================================

def find_images(directory: str, extensions: List[str] = None) -> List[str]:
    """Find all images in directory (case‑insensitive, no duplicates)"""
    if extensions is None:
        extensions = ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.webp']
    
    images = []
    for f in Path(directory).iterdir():
        if f.is_file() and f.suffix.lower() in extensions:
            images.append(str(f))
    return sorted(images)

def main():
    parser = argparse.ArgumentParser(
        description='Comprehensive OCR tool using Tesseract',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Single image to Word
  python ocr_script.py image.jpg -o output.docx --format docx --preprocess
  
  # Batch process all images in folder to Word docs (outputs: image1.docx, image2.docx...)
  python ocr_script.py -d images/ -o results/ --format docx --preprocess
        """
    )
    
    # Input options
    input_group = parser.add_mutually_exclusive_group(required=True)
    input_group.add_argument('image', nargs='?', help='Input image file path (single mode)')
    input_group.add_argument('-d', '--directory', help='Directory containing images for batch processing')
    
    # Output options
    parser.add_argument('-o', '--output', required=True, help='Output file (single mode) or directory (batch mode)')
    parser.add_argument('--format', choices=['txt', 'json', 'csv', 'hocr', 'pdf', 'docx'], 
                       default='txt', help='Output format (default: txt)')
    
    # OCR options
    parser.add_argument('-l', '--lang', default='eng', help='OCR language (default: eng)')
    parser.add_argument('--psm', type=int, default=6, choices=[0,1,3,4,5,6,7,8,9,11,12,13],
                       help='Page segmentation mode (default: 6)')
    parser.add_argument('--confidence', type=float, default=0.0,
                       help='Minimum confidence threshold (0-100)')
    parser.add_argument('--tesseract-path', help='Path to tesseract executable')
    parser.add_argument('--tessdata-dir', help='Path to tessdata directory')
    
    # Preprocessing options
    parser.add_argument('--preprocess', action='store_true', help='Enable image preprocessing')
    parser.add_argument('--no-threshold', action='store_true', help='Disable thresholding')
    parser.add_argument('--no-deskew', action='store_true', help='Disable deskewing')
    parser.add_argument('--threshold-value', type=int, default=150, help='Threshold value (default: 150)')
    parser.add_argument('--resize', action='store_true', help='Upscale image for better OCR')
    parser.add_argument('--resize-factor', type=float, default=2.0, help='Resize factor (default: 2.0)')
    parser.add_argument('--sharpen', action='store_true', help='Apply sharpening filter')
    parser.add_argument('--save-debug', action='store_true', help='Save preprocessed debug images')
    
    # Performance options
    parser.add_argument('--workers', type=int, default=4, help='Number of parallel workers (default: 4)')
    parser.add_argument('--verbose', action='store_true', help='Enable verbose output')
    
    args = parser.parse_args()
    
    # Build configurations
    psm_mode = PSMMode(args.psm) if args.psm in [m.value for m in PSMMode] else PSMMode.UNIFORM_BLOCK
    oem_mode = OEMMode.DEFAULT
    
    ocr_config = OCRConfig(
        language=args.lang,
        psm=psm_mode,
        oem=oem_mode,
        tessdata_dir=args.tessdata_dir,
        tesseract_path=args.tesseract_path,
        confidence_threshold=args.confidence,
        output_format=args.format
    )
    
    preprocess_config = PreprocessConfig(
        grayscale=True,
        threshold=not args.no_threshold if args.preprocess else False,
        threshold_value=args.threshold_value,
        denoise=args.preprocess,
        resize=args.resize,
        resize_factor=args.resize_factor,
        sharpen=args.sharpen,
        contrast_enhance=args.preprocess,
        deskew=not args.no_deskew if args.preprocess else False,
        remove_borders=args.preprocess
    )
    
    # Initialize components
    preprocessor = ImagePreprocessor(preprocess_config)
    ocr_engine = OCREngine(ocr_config)
    output_handler = OutputHandler()
    
    # Process based on mode
    if args.directory:
        # Batch mode
        image_dir = args.directory
        if not os.path.isdir(image_dir):
            print(f"Error: Directory not found: {image_dir}")
            sys.exit(1)
        
        images = find_images(image_dir)
        if not images:
            print(f"No images found in {image_dir}")
            sys.exit(1)
        
        print(f"Found {len(images)} images in '{image_dir}'")
        
        output_dir = args.output
        os.makedirs(output_dir, exist_ok=True)
        
        processor = BatchOCRProcessor(ocr_engine, preprocessor, output_handler, args.workers)
        results = processor.process_batch(images, output_dir, args.format, args.save_debug)
        
        # Summary
        successful = sum(1 for r in results if r['success'])
        print(f"\n{'='*50}")
        print(f"Batch processing complete: {successful}/{len(results)} successful")
        print(f"Output directory: {output_dir}")
        
        # Save combined results if CSV format
        if args.format == 'csv':
            combined_path = os.path.join(output_dir, 'combined_results.csv')
            output_handler.save_csv(results, combined_path)
            print(f"Saved combined results to {combined_path}")
        
    elif args.image:
        # Single image mode
        if not os.path.exists(args.image):
            print(f"Error: Image not found: {args.image}")
            sys.exit(1)
        
        print(f"Processing: {args.image}")
        
        # Preprocess
        if args.preprocess:
            print("Preprocessing image...")
            processed = preprocessor.preprocess(args.image)
            if args.save_debug:
                debug_path = f"debug_{Path(args.image).stem}.png"
                preprocessor.save_debug_image(processed, debug_path)
                print(f"Saved debug image: {debug_path}")
        else:
            # Just read the image without preprocessing
            processed = cv2.imread(args.image)
            if processed is None:
                print(f"Error: Could not read image {args.image}")
                sys.exit(1)
            if len(processed.shape) == 3:
                processed = cv2.cvtColor(processed, cv2.COLOR_BGR2GRAY)
        
        # Extract text
        print("Running OCR...")
        text, metadata = ocr_engine.extract_text(processed)
        
        # Save output
        output_path = args.output
        # If output is a directory (should not happen, but handle gracefully)
        if os.path.isdir(output_path):
            base_name = Path(args.image).stem
            output_path = os.path.join(output_path, f"{base_name}.{args.format}")
        
        if args.format == 'txt':
            output_handler.save_txt(text, output_path, metadata)
        elif args.format == 'docx':
            output_handler.save_docx(text, output_path, metadata, args.image)
        elif args.format == 'json':
            output_handler.save_json(text, output_path, metadata, args.image)
        elif args.format == 'hocr':
            hocr_content = ocr_engine.get_hocr(processed)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(hocr_content)
        elif args.format == 'pdf':
            pdf_content = ocr_engine.get_pdf(processed)
            with open(output_path, 'wb') as f:
                f.write(pdf_content)
        
        print(f"✓ Saved to: {output_path}")
        print(f"  Word count: {metadata.get('word_count', 0)}")
        print(f"  Avg confidence: {metadata.get('avg_confidence', 0):.2f}%")
        
        if args.verbose:
            print(f"\n{'='*50}")
            print("Extracted Text:")
            print(f"{'='*50}")
            print(text)
    
    else:
        # This should never happen because args are mutually exclusive required
        parser.print_help()
        sys.exit(1)

if __name__ == '__main__':
    main()